"""
Generate a Word Document for AI Expense Tracker Project
Run: python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Calibri'
    return p

def add_para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        # before bold
        if idx > 0:
            r1 = p.add_run(text[:idx])
            set_font(r1)
        r2 = p.add_run(bold_part)
        set_font(r2, bold=True)
        rest = text[idx+len(bold_part):]
        if rest:
            r3 = p.add_run(rest)
            set_font(r3)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=10, color=(255,255,255))
        # Dark blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = 'DEEAF1' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    # Column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    return table

def add_code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 100, 0)

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('AI Expense Tracker')
r.font.name  = 'Calibri'
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(31, 73, 125)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('Project Documentation')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(16)
r2.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info.add_run(f'Generated on: {datetime.datetime.now().strftime("%B %d, %Y")}')
r3.font.name  = 'Calibri'
r3.font.size  = Pt(12)
r3.font.color.rgb = RGBColor(127, 127, 127)

doc.add_paragraph()
doc.add_paragraph()

# Divider line
div = doc.add_paragraph('─' * 60)
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in div.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. Project Overview', level=1)

add_para(
    'The AI Expense Tracker is a full-stack personal finance management web application '
    'built using Python (Flask) and MongoDB. It allows users to securely register, log in, '
    'and track their income and expenses with real-time visual analytics, budget limit alerts, '
    'and AI-powered spending suggestions.'
)

doc.add_paragraph()
add_heading('1.1 Project Goals', level=2)
add_bullet('Track personal income and expenses with category breakdown.')
add_bullet('Visualize spending using interactive charts (Bar, Pie, Doughnut, Line, Polar Area, Radar).')
add_bullet('Set per-category spending limits and receive real-time budget alerts.')
add_bullet('Provide AI-based spending tips based on user behavior.')
add_bullet('Ensure user data isolation with session-based multi-user support.')

doc.add_paragraph()
add_heading('1.2 Key Highlights', level=2)
add_bullet('Multi-user secure login and registration system.')
add_bullet('MongoDB NoSQL database for flexible data storage.')
add_bullet('Premium dark-themed responsive UI with smooth animations.')
add_bullet('Interactive chart switcher with 6 chart types on the dashboard.')
add_bullet('Budget alert notifications (toast pop-ups + inline banners).')
add_bullet('Custom animated delete confirmation modal.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Technology Stack', level=1)

add_heading('2.1 Backend', level=2)
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Web Framework', 'Flask (Python)', 'Handles HTTP routes, sessions, and template rendering'],
        ['Database', 'MongoDB (NoSQL)', 'Stores users, expenses, income, and category limits'],
        ['DB Driver', 'PyMongo', 'Python library to connect Flask with MongoDB'],
        ['DB Server', 'localhost:27017', 'Local MongoDB instance running on the machine'],
        ['Session Mgmt', 'Flask Sessions', 'Keeps users logged in using secret key + cookies'],
    ],
    col_widths=[1.5, 1.8, 3.5]
)

doc.add_paragraph()
add_heading('2.2 Frontend', level=2)
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Markup', 'HTML5 + Jinja2', 'Server-side templating rendered by Flask'],
        ['Styling', 'Vanilla CSS', 'Custom CSS with variables, animations, glassmorphism'],
        ['Charts', 'Chart.js (CDN)', 'Interactive data visualization (6 chart types)'],
        ['Fonts', 'Google Fonts – Inter', 'Modern typography across all pages'],
        ['JavaScript', 'Vanilla JS', 'Chart switcher, modals, toast notifications'],
    ],
    col_widths=[1.5, 1.8, 3.5]
)

doc.add_paragraph()
add_heading('2.3 No Frameworks Used', level=2)
add_para('The following are intentionally NOT used in this project:')
add_bullet('❌ No Bootstrap or Tailwind CSS')
add_bullet('❌ No React, Vue, or Angular')
add_bullet('❌ No jQuery')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. PROJECT STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Project Structure', level=1)

add_code_block(
    'AI_Expense_Tracker/\n'
    '├── app.py                  # Main Flask application (routes & logic)\n'
    '├── templates/\n'
    '│   ├── base_sidebar.html   # Base layout with sidebar, toast system, theme toggle\n'
    '│   ├── base_auth.html      # Base layout for login/register pages\n'
    '│   ├── index.html          # Landing / home page\n'
    '│   ├── login.html          # Login form\n'
    '│   ├── register.html       # Registration form\n'
    '│   ├── dashboard.html      # Main dashboard with charts & analytics\n'
    '│   ├── add_expense.html    # Add new expense form\n'
    '│   ├── add_income.html     # Add new income form\n'
    '│   ├── view_expenses.html  # Table of all expenses with edit/delete\n'
    '│   ├── edit_expense.html   # Edit an existing expense\n'
    '│   ├── reports.html        # Monthly/yearly financial reports\n'
    '│   ├── success.html        # Registration success page\n'
    '│   └── error.html          # Error / login failed page\n'
    '└── venv/                   # Python virtual environment'
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. DATABASE DESIGN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. Database Design', level=1)
add_para('Database Name: expense_tracker  |  Type: MongoDB (NoSQL)')
doc.add_paragraph()

add_heading('4.1 Collections', level=2)

add_heading('users', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['_id', 'ObjectId', 'Auto-generated unique identifier'],
        ['username', 'String', 'Unique username for login'],
        ['password', 'String', 'User password (plain text – upgrade recommended)'],
    ],
    col_widths=[1.5, 1.5, 3.8]
)

doc.add_paragraph()
add_heading('expenses', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['_id', 'ObjectId', 'Auto-generated unique identifier'],
        ['user_id', 'String', 'References the logged-in user (_id as string)'],
        ['amount', 'Integer', 'Expense amount in Indian Rupees (₹)'],
        ['category', 'String', 'Category: Food, Travel, Shopping, Bills, Entertainment, Other'],
        ['date', 'String', 'Date in YYYY-MM-DD format'],
        ['note', 'String', 'Optional description/note for the expense'],
    ],
    col_widths=[1.5, 1.2, 4.1]
)

doc.add_paragraph()
add_heading('income', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['_id', 'ObjectId', 'Auto-generated unique identifier'],
        ['user_id', 'String', 'References the logged-in user'],
        ['amount', 'Integer', 'Income amount in Indian Rupees (₹)'],
        ['source', 'String', 'Income source (e.g., Salary, Freelance, etc.)'],
        ['date', 'String', 'Date in YYYY-MM-DD format'],
        ['note', 'String', 'Optional note'],
    ],
    col_widths=[1.5, 1.2, 4.1]
)

doc.add_paragraph()
add_heading('category_limits', level=3)
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['_id', 'ObjectId', 'Auto-generated unique identifier'],
        ['user_id', 'String', 'References the logged-in user'],
        ['limits', 'Object', 'Key-value pairs: { "Food": 2000, "Shopping": 5000, ... }'],
    ],
    col_widths=[1.5, 1.2, 4.1]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. APPLICATION FEATURES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. Application Features', level=1)

add_heading('5.1 User Authentication', level=2)
add_bullet('User Registration: New users can register with a username and password.')
add_bullet('User Login: Authenticated via username + password lookup in MongoDB.')
add_bullet('Session Management: Flask sessions store user_id and username after login.')
add_bullet('Logout: Session is cleared and user is redirected to login.')
add_bullet('Data Isolation: All queries filter by session user_id to prevent cross-user data access.')

doc.add_paragraph()
add_heading('5.2 Expense Management', level=2)
add_bullet('Add Expense: Form to input amount, category, date, and optional note.')
add_bullet('View Expenses: Full table showing all expenses with Edit and Delete actions.')
add_bullet('Edit Expense: Pre-filled form to update amount, category, and date.')
add_bullet('Delete Expense: Custom animated modal confirmation before deletion.')

doc.add_paragraph()
add_heading('5.3 Income Management', level=2)
add_bullet('Add Income: Form to input amount, source, date, and optional note.')
add_bullet('Income is aggregated to calculate Net Balance (Income – Expenses).')

doc.add_paragraph()
add_heading('5.4 Dashboard & Analytics', level=2)
add_bullet('Stat Cards: Total Income, Total Spent, Net Balance, Top Spending Category.')
add_bullet('Interactive Charts: One chart card with 6 switchable chart types:')
add_bullet('  → Bar Chart – Compare categories side by side.')
add_bullet('  → Pie Chart – Proportional breakdown of spending.')
add_bullet('  → Doughnut Chart – Clean ring chart with legend.')
add_bullet('  → Line Chart – Trend of spending over time.')
add_bullet('  → Polar Area – Mixed size area comparison.')
add_bullet('  → Radar Chart – Multi-axis spending overview.')
add_bullet('Period Filter: Charts can be filtered by All Time, specific Month, or Year.')
add_bullet('Category Summary Pills: Shows ₹ amount and % for each category below chart.')
add_bullet('Recent Expenses: Table showing last 5 expenses with quick edit link.')
add_bullet('AI Suggestions: Spending tips shown after adding an expense.')

doc.add_paragraph()
add_heading('5.5 Budget Alert Notifications', level=2)
add_bullet('Category Spending Limits: Users can set a ₹ limit per category.')
add_bullet('Toast Notifications: Animated pop-ups appear on page load for exceeded/near-limit categories.')
add_bullet('Inline Alert Banner: Grouped alert panel shown at top of dashboard.')
add_bullet('Sidebar Bell (🔔): Displays the count of active alerts with a pulsing badge.')
add_bullet('Progress Bars: Visual bars inside each limit card showing % used (green/yellow/red).')
add_bullet('Alert Levels: 🚨 Exceeded (>100%) and ⚠️ Warning (≥80%).')

doc.add_paragraph()
add_heading('5.6 Reports Page', level=2)
add_bullet('Monthly totals for expenses and income for the current month.')
add_bullet('Yearly totals for expenses and income for the current year.')
add_bullet('Category breakdown table with total spent, count, and percentage.')

doc.add_paragraph()
add_heading('5.7 UI/UX Features', level=2)
add_bullet('Dark/Light Mode Toggle: Persistent theme switcher using localStorage.')
add_bullet('Premium dark theme with CSS variables, gradients, and glassmorphism.')
add_bullet('Smooth hover animations and micro-interactions on all interactive elements.')
add_bullet('Category color pills (Food=pink, Travel=purple, Shopping=orange, etc.).')
add_bullet('Custom delete modal with backdrop blur instead of browser confirm() dialog.')
add_bullet('Auto-fading AI suggestion alerts after 4 seconds.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6. Flask Application Routes', level=1)

add_table(
    ['Route', 'Method', 'Description'],
    [
        ['/', 'GET', 'Home / Landing page'],
        ['/register', 'GET, POST', 'User registration form and handler'],
        ['/login', 'GET, POST', 'User login form and handler'],
        ['/logout', 'GET', 'Clears session and redirects to login'],
        ['/dashboard', 'GET', 'Main dashboard with stats and charts'],
        ['/add-expense-page', 'GET', 'Shows Add Expense form'],
        ['/add-expense', 'POST', 'Saves new expense to MongoDB'],
        ['/edit-expense/<id>', 'GET', 'Shows Edit Expense form pre-filled'],
        ['/edit-expense/<id>', 'POST', 'Updates expense in MongoDB'],
        ['/delete-expense/<id>', 'GET', 'Deletes expense from MongoDB'],
        ['/add-income-page', 'GET', 'Shows Add Income form'],
        ['/add-income', 'POST', 'Saves new income to MongoDB'],
        ['/view-expenses', 'GET', 'Shows all expenses in a table'],
        ['/reports', 'GET', 'Shows monthly/yearly financial reports'],
        ['/set-category-limit', 'POST', 'Sets a spending limit for a category'],
        ['/remove-category-limit/<cat>', 'GET', 'Removes a category spending limit'],
    ],
    col_widths=[2.3, 1.2, 3.3]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7. How to Run the Project', level=1)

add_heading('7.1 Prerequisites', level=2)
add_bullet('Python 3.8 or higher installed.')
add_bullet('MongoDB installed and running on localhost:27017.')
add_bullet('pip (Python package manager).')

doc.add_paragraph()
add_heading('7.2 Installation Steps', level=2)

steps = [
    ('Step 1 – Clone or download the project folder.', None),
    ('Step 2 – Create and activate a virtual environment:', 'python -m venv venv\nvenv\\Scripts\\activate   # Windows'),
    ('Step 3 – Install required packages:', 'pip install flask pymongo'),
    ('Step 4 – Start MongoDB service on your machine.', None),
    ('Step 5 – Run the Flask application:', 'python app.py'),
    ('Step 6 – Open your browser and go to:', 'http://127.0.0.1:5000'),
]

for label, code in steps:
    add_para(label, bold=True, size=11)
    if code:
        add_code_block(code)
    doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. FUTURE ENHANCEMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8. Future Enhancements', level=1)

add_table(
    ['Feature', 'Description', 'Priority'],
    [
        ['Password Hashing', 'Hash passwords using bcrypt instead of plain text storage', 'HIGH'],
        ['Export to CSV/PDF', 'Allow users to download their expense reports', 'HIGH'],
        ['Savings Goals', 'Set a savings target and track progress visually', 'MEDIUM'],
        ['Recurring Expenses', 'Mark bills/EMIs as recurring with monthly reminders', 'MEDIUM'],
        ['Search & Filter', 'Filter expenses by category, date range, or keyword', 'MEDIUM'],
        ['Edit Income', 'Add edit and delete functionality for income records', 'MEDIUM'],
        ['Mobile Responsive', 'Collapsible sidebar for mobile/tablet screens', 'MEDIUM'],
        ['Email Reports', 'Weekly/monthly summary sent via email (Flask-Mail)', 'LOW'],
        ['AI Prediction', 'Predict next month spending using historical patterns', 'LOW'],
        ['Multi-currency', 'Support USD, EUR with live conversion rates', 'LOW'],
    ],
    col_widths=[1.8, 3.8, 1.2]
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('9. Project Summary', level=1)

add_para(
    'The AI Expense Tracker is a complete, production-ready personal finance web application '
    'built with modern web technologies. It successfully demonstrates the integration of a '
    'Python Flask backend with a MongoDB NoSQL database, a Jinja2 templating system, and '
    'a fully hand-crafted premium frontend UI.',
    size=11
)
doc.add_paragraph()
add_para(
    'The application supports real-time expense and income tracking, interactive multi-type '
    'chart analytics, smart budget alert notifications, and a polished user experience with '
    'dark/light mode, animated modals, and toast notifications — all without any CSS or JS frameworks.',
    size=11
)
doc.add_paragraph()
add_para(
    'This project can serve as a strong portfolio piece demonstrating full-stack development '
    'skills, database design, REST API design, UI/UX thinking, and clean code organization.',
    size=11
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\hema\OneDrive\Desktop\AI_Expense_Tracker\AI_Expense_Tracker_Documentation.docx'
doc.save(output_path)
print(f"\n✅ Word document saved successfully!\n📄 Location: {output_path}\n")
